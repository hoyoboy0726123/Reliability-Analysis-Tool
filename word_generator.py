"""
Word 測試報告生成器
使用 python-docx 生成專業的可靠度測試報告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io
import re
import base64
from PIL import Image as PILImage

# 導入圖表生成模組
try:
    from chart_generator import generate_all_charts
    CHARTS_AVAILABLE = True
except ImportError as e:
    print(f"Warning: chart_generator not available: {e}")
    CHARTS_AVAILABLE = False


def format_af_value(value):
    """安全格式化 AF 值"""
    if value is None or value == 'N/A' or (isinstance(value, str) and value.strip() == ''):
        return 'N/A'
    try:
        return f"{float(value):,.4f}"
    except (ValueError, TypeError):
        return str(value)


def add_colored_text(paragraph, text, color=None, bold=False, size=None):
    """添加帶顏色和樣式的文字到段落"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if color:
        if isinstance(color, str):
            # 處理十六進制顏色
            color = color.lstrip('#')
            rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
        elif isinstance(color, tuple):
            run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return run


def set_cell_background(cell, color):
    """設置表格單元格背景顏色"""
    if isinstance(color, str):
        color = color.lstrip('#')

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def parse_html_conclusion(doc, html_text):
    """
    將 HTML 格式的結論轉換為 Word 段落
    使用與 PDF 相同的處理邏輯
    """
    # 替換特殊符號（與 PDF 一致）
    html_text = html_text.replace('⚠️', '[!]').replace('✓', '[OK]').replace('✗', '[X]')

    # 轉換 HTML 標籤為統一格式（與 PDF 一致）
    html_text = re.sub(r'<strong class="text-warning">(.*?)</strong>',
                       r'<font color="#f59e0b"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-danger">(.*?)</strong>',
                       r'<font color="#dc2626"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-success">(.*?)</strong>',
                       r'<font color="#10b981"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-info">(.*?)</strong>',
                       r'<font color="#0ea5e9"><b>\1</b></font>', html_text, flags=re.DOTALL)

    # 處理一般 strong 標籤（與 PDF 一致）
    html_text = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', html_text, flags=re.DOTALL)

    # 處理換行和列表（與 PDF 一致）
    html_text = html_text.replace('<br><br>', '|||PARA|||').replace('<br>', '<br/>')
    html_text = html_text.replace('<ul class="mb-0 mt-2">', '').replace('<ul>', '').replace('</ul>', '')
    html_text = re.sub(r'<li>(.*?)</li>', r'• \1<br/>', html_text, flags=re.DOTALL)
    html_text = re.sub(r' class="[^"]*"', '', html_text)

    # 分段處理（與 PDF 一致）
    paragraphs = html_text.split('|||PARA|||')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 移除空白的 div 標籤
        para = re.sub(r'<div[^>]*>\s*</div>', '', para)
        para = para.replace('<div class="small text-muted mt-2">', '<font size="9" color="#6b7280">')
        para = para.replace('</div>', '</font>')

        try:
            # 創建段落並解析內容
            p = doc.add_paragraph()
            parse_paragraph_content(p, para)
        except Exception as e:
            # 如果失敗，使用純文字版本
            print(f"Error rendering conclusion paragraph: {e}")
            plain_text = re.sub(r'<[^>]+>', '', para)
            p = doc.add_paragraph(plain_text)


def parse_paragraph_content(paragraph, html_text):
    """
    解析段落內容，處理格式標籤
    支援：<font color="..."><b>text</b></font>, <b>text</b>, <br/>
    """
    # 處理帶顏色的粗體文字：<font color="#xxx"><b>text</b></font>
    pattern_colored_bold = r'<font color="([^"]+)"><b>(.*?)</b></font>'

    # 處理純粗體文字：<b>text</b>
    pattern_bold = r'<b>(.*?)</b>'

    # 處理換行：<br/>
    html_text = html_text.replace('<br/>', '\n')

    pos = 0
    while pos < len(html_text):
        # 尋找下一個標籤
        match_colored = re.search(pattern_colored_bold, html_text[pos:], re.DOTALL)
        match_bold = re.search(pattern_bold, html_text[pos:], re.DOTALL)

        # 確定最近的匹配
        next_match = None
        match_type = None

        if match_colored and match_bold:
            if match_colored.start() < match_bold.start():
                next_match = match_colored
                match_type = 'colored'
            else:
                next_match = match_bold
                match_type = 'bold'
        elif match_colored:
            next_match = match_colored
            match_type = 'colored'
        elif match_bold:
            next_match = match_bold
            match_type = 'bold'

        if next_match:
            # 添加匹配前的普通文字
            if next_match.start() > 0:
                plain_text = html_text[pos:pos + next_match.start()]
                if plain_text:
                    paragraph.add_run(plain_text)

            # 添加格式化文字
            if match_type == 'colored':
                color = next_match.group(1)
                text = next_match.group(2)
                add_colored_text(paragraph, text, color=color, bold=True)
            else:  # bold
                text = next_match.group(1)
                run = paragraph.add_run(text)
                run.bold = True

            pos = pos + next_match.end()
        else:
            # 沒有更多標籤，添加剩餘文字
            remaining = html_text[pos:]
            if remaining:
                paragraph.add_run(remaining)
            break


def generate_word_report(data, output_file):
    """
    生成可靠度測試報告 Word 文件

    Args:
        data: 包含所有測試數據和結果的字典
        output_file: 輸出文件對象 (BytesIO)
    """
    # 創建文件
    doc = Document()

    # 設定默認字體（支持中文）
    doc.styles['Normal'].font.name = 'Microsoft JhengHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # === 封面頁 ===
    # 標題
    title = doc.add_heading('可靠度測試報告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = RGBColor(26, 84, 144)  # 藍色

    doc.add_paragraph()  # 空行

    # 副標題
    subtitle = doc.add_paragraph('Reliability Test Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # 報告信息
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light List Accent 1'

    info_data = [
        ('報告日期', datetime.now().strftime('%Y-%m-%d')),
        ('報告時間', datetime.now().strftime('%H:%M:%S')),
        ('分析模式', '韋伯分析' if data.get('analysis_mode') == 'weibull' else '零失效分析')
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value

    # 分頁
    doc.add_page_break()

    # === 第1節：加速因子計算 ===
    heading = doc.add_heading('1. 加速因子計算結果', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

    af_result = data.get('results', {}).get('af_result', {})

    # AF 總值
    af_total = af_result.get('af_total', 'N/A')
    p = doc.add_paragraph()
    p.add_run('總加速因子 (AF Total): ').bold = True
    add_colored_text(p, format_af_value(af_total), color='#d97706', bold=True, size=14)

    doc.add_paragraph()

    # AF 分項表格
    af_params = data.get('af_params', {})
    table_data = []

    # 溫度
    if af_params.get('enable_temp', False):
        table_data.append([
            '溫度加速 (Temperature)',
            format_af_value(af_result.get('af_t', 'N/A')),
            f"T_use={af_params.get('t_use', 'N/A')}°C, T_alt={af_params.get('t_alt', 'N/A')}°C, Ea={af_params.get('ea', 'N/A')} eV"
        ])

    # 濕度
    if af_params.get('enable_hum', False):
        table_data.append([
            '濕度加速 (Humidity)',
            format_af_value(af_result.get('af_rh', 'N/A')),
            f"RH_use={af_params.get('rh_use', 'N/A')}%, RH_alt={af_params.get('rh_alt', 'N/A')}%, n={af_params.get('n_hum', 'N/A')}"
        ])

    # 電壓
    if af_params.get('enable_voltage', False):
        table_data.append([
            '電壓加速 (Voltage)',
            format_af_value(af_result.get('af_v', 'N/A')),
            f"V_use={af_params.get('v_use', 'N/A')}V, V_alt={af_params.get('v_alt', 'N/A')}V, β={af_params.get('beta_v', 'N/A')}"
        ])

    # 熱循環
    if af_params.get('enable_tc', False):
        table_data.append([
            '熱循環加速 (Thermal Cycling)',
            format_af_value(af_result.get('af_tc', 'N/A')),
            f"ΔT_use={af_params.get('dt_use', 'N/A')}°C, ΔT_alt={af_params.get('dt_alt', 'N/A')}°C"
        ])

    # 振動
    if af_params.get('enable_vib', False):
        table_data.append([
            '振動加速 (Vibration)',
            format_af_value(af_result.get('af_vib', 'N/A')),
            f"G_use={af_params.get('g_use', 'N/A')}g, G_alt={af_params.get('g_alt', 'N/A')}g, n={af_params.get('n_vib', 'N/A')}"
        ])

    # UV
    if af_params.get('enable_uv', False):
        table_data.append([
            'UV 輻射加速 (UV Radiation)',
            format_af_value(af_result.get('af_uv', 'N/A')),
            f"t_field={af_params.get('t_field_uv', 'N/A')}h, t_accel={af_params.get('t_accel_uv', 'N/A')}h"
        ])

    # 化學
    if af_params.get('enable_chem', False):
        table_data.append([
            '化學濃度加速 (Chemical)',
            format_af_value(af_result.get('af_chem', 'N/A')),
            f"C_use={af_params.get('c_use', 'N/A')}, C_alt={af_params.get('c_alt', 'N/A')}, n={af_params.get('n_chem', 'N/A')}"
        ])

    if table_data:
        table = doc.add_table(rows=len(table_data) + 1, cols=3)
        table.style = 'Light Grid Accent 1'

        # 表頭
        header_cells = table.rows[0].cells
        header_cells[0].text = '加速因子類型'
        header_cells[1].text = 'AF 值'
        header_cells[2].text = '參數條件'

        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_background(cell, '1a5490')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # 數據行
        for i, row_data in enumerate(table_data, start=1):
            cells = table.rows[i].cells
            for j, value in enumerate(row_data):
                cells[j].text = str(value)

    doc.add_page_break()

    # === 第2節：測試配置 ===
    heading = doc.add_heading('2. 測試配置', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

    test_data = data.get('test_data', {})

    config_table = doc.add_table(rows=4, cols=2)
    config_table.style = 'Light List Accent 1'

    config_data = [
        ('樣品數量', test_data.get('n_samples', 'N/A')),
        ('測試時間', f"{test_data.get('t_test', 'N/A')} 小時"),
        ('信心水準', f"{float(test_data.get('cl', 0.6)) * 100:.0f}%"),
        ('任務時間', f"{test_data.get('mission_years', 'N/A')} 年")
    ]

    for i, (label, value) in enumerate(config_data):
        config_table.rows[i].cells[0].text = label
        config_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        config_table.rows[i].cells[1].text = str(value)

    # 失效數據
    failures_str = test_data.get('failures', '')
    if failures_str and failures_str.strip():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('失效時間數據: ').bold = True
        p.add_run(failures_str)

    doc.add_page_break()

    # === 第3節：韋伯分析結果 (如果有) ===
    weibull_result = data.get('results', {}).get('weibull_result', {})
    if weibull_result and 'beta' in weibull_result:
        heading = doc.add_heading('3. 韋伯分析結果', level=1)
        heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

        # 韋伯參數
        weibull_table = doc.add_table(rows=4, cols=2)
        weibull_table.style = 'Light List Accent 1'

        weibull_data = [
            ('形狀參數 (β)', f"{weibull_result.get('beta', 'N/A'):.4f}"),
            ('特性壽命 - ALT (η_alt)', f"{weibull_result.get('eta_alt', 'N/A'):.2f} 小時"),
            ('擬合優度 (R²)', f"{weibull_result.get('r_squared', 'N/A'):.4f}"),
            ('分析方法', weibull_result.get('method', 'BENARD + RRY'))
        ]

        for i, (label, value) in enumerate(weibull_data):
            weibull_table.rows[i].cells[0].text = label
            weibull_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            weibull_table.rows[i].cells[1].text = str(value)

        doc.add_page_break()

    # === 第4節：可靠度結果 ===
    heading = doc.add_heading('4. 可靠度評估結果', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

    reliability_result = data.get('results', {}).get('reliability_result', {})

    # 韋伯模式結果
    if 'weibull' in reliability_result:
        weibull_rel = reliability_result['weibull']

        doc.add_heading('韋伯分析模式', level=2)

        rel_table = doc.add_table(rows=5, cols=2)
        rel_table.style = 'Light List Accent 1'

        bx_percent = weibull_rel.get('bx_percent', 1)

        rel_data = [
            ('特性壽命 - 現場 (η_use)', f"{weibull_rel.get('eta_use', 'N/A'):,.2f} 小時"),
            ('平均壽命 (MTTF)', f"{weibull_rel.get('mttf_use', 'N/A'):,.2f} 小時 ({weibull_rel.get('mttf_use', 0) / 8760:.2f} 年)"),
            ('任務可靠度 (R_mission)', f"{weibull_rel.get('r_mission', 'N/A') * 100:.2f}%"),
            (f'B{bx_percent}% 壽命', f"{weibull_rel.get('bx_life', 'N/A'):,.2f} 小時 ({weibull_rel.get('bx_life', 0) / 8760:.2f} 年)"),
            ('失效風險', f"{(1 - weibull_rel.get('r_mission', 0)) * 100:.2f}%")
        ]

        for i, (label, value) in enumerate(rel_data):
            rel_table.rows[i].cells[0].text = label
            rel_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            rel_table.rows[i].cells[1].text = str(value)

    # 零失效模式結果
    if 'zero_failure' in reliability_result:
        zf_rel = reliability_result['zero_failure']

        doc.add_paragraph()
        doc.add_heading('零失效分析模式', level=2)

        zf_table = doc.add_table(rows=4, cols=2)
        zf_table.style = 'Light List Accent 1'

        zf_data = [
            ('MTTF 下限 - 現場', f"{zf_rel.get('mttf_use_lower', 'N/A'):,.2f} 小時"),
            ('失效率上限 - 現場', f"{zf_rel.get('lambda_use_upper', 'N/A'):,.2f} FITs"),
            ('任務可靠度 (R_mission)', f"{zf_rel.get('r_mission', 'N/A') * 100:.2f}%"),
            ('失效風險', f"{(1 - zf_rel.get('r_mission', 0)) * 100:.2f}%")
        ]

        for i, (label, value) in enumerate(zf_data):
            zf_table.rows[i].cells[0].text = label
            zf_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            zf_table.rows[i].cells[1].text = str(value)

    # === 第5節：圖表 (如果有) ===
    charts_added = False
    if CHARTS_AVAILABLE:
        try:
            # 提取數據
            analysis_mode = data.get('analysis_mode', 'weibull')
            weibull_result = data.get('results', {}).get('weibull_result')
            reliability_result = data.get('results', {}).get('reliability_result')

            # 嘗試生成圖表
            charts = generate_all_charts(
                analysis_mode=analysis_mode,
                weibull_result=weibull_result if analysis_mode == 'weibull' else None,
                reliability_result=reliability_result,
                max_time=50000
            )

            # 檢查是否有有效的圖表
            valid_charts = {k: v for k, v in charts.items() if v}

            if valid_charts:
                doc.add_page_break()
                heading = doc.add_heading('5. 分析圖表', level=1)
                heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

                # 添加圖表
                for i, (chart_name, chart_buffer) in enumerate(valid_charts.items(), start=1):
                    try:
                        # chart_buffer 是 BytesIO 對象
                        chart_buffer.seek(0)  # 重置指針到開始

                        # 添加圖表標題
                        chart_titles = {
                            'reliability': '可靠度曲線 R(t)',
                            'failure_rate': '失效率曲線 λ(t)',
                            'pdf': '機率密度函數 f(t)'
                        }

                        doc.add_heading(f'5.{i} {chart_titles.get(chart_name, chart_name)}', level=2)

                        # 插入圖片（直接使用 BytesIO）
                        doc.add_picture(chart_buffer, width=Inches(6))

                        # 置中
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        doc.add_paragraph()
                        charts_added = True
                    except Exception as e:
                        print(f"Error adding chart {chart_name}: {e}")
                        import traceback
                        traceback.print_exc()

        except Exception as e:
            print(f"Error generating charts for Word: {e}")
            import traceback
            traceback.print_exc()

    # === 結論節 ===
    doc.add_page_break()
    section_num = 6 if charts_added else 5
    heading = doc.add_heading(f'{section_num}. 測試結論', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 84, 144)

    conclusion_html = data.get('conclusion', '')
    if conclusion_html and conclusion_html.strip():
        # 使用新的解析器處理結論
        parse_html_conclusion(doc, conclusion_html)
    else:
        doc.add_paragraph('無結論內容')

    # 保存文件
    doc.save(output_file)


def generate_report_from_request(data):
    """
    從 Flask 請求生成 Word 報告

    Args:
        data: 來自前端的請求數據字典

    Returns:
        BytesIO: Word 文件的內存緩衝區
    """
    try:
        # 創建內存緩衝區
        buffer = io.BytesIO()

        # 生成 Word 報告
        generate_word_report(data, buffer)

        # 重置緩衝區位置
        buffer.seek(0)

        return buffer

    except Exception as e:
        print(f"Error in generate_report_from_request: {e}")
        import traceback
        traceback.print_exc()
        raise
