"""
Word 测试报告生成器 V2 - 与 PDF 格式完全一致
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io
import re

def format_af_value(value):
    """安全格式化 AF 值"""
    if value is None or value == 'N/A' or (isinstance(value, str) and value.strip() == ''):
        return 'N/A'
    try:
        return f"{float(value):,.4f}"
    except (ValueError, TypeError):
        return str(value)


def set_cell_background(cell, color):
    """设置表格单元格背景颜色"""
    if isinstance(color, str):
        color = color.lstrip('#')

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, color='d1d5db', size='6'):
    """设置表格边框"""
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_colored_text(paragraph, text, color=None, bold=False, size=None):
    """添加带颜色和样式的文字"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if color:
        if isinstance(color, str):
            color = color.lstrip('#')
            rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
    if size:
        run.font.size = Pt(size)
    return run


def parse_html_conclusion(doc, html_text):
    """解析 HTML 格式的结论 - 与 PDF 完全一致的逻辑"""
    # 替换特殊符号
    html_text = html_text.replace('⚠️', '[!]').replace('✓', '[OK]').replace('✗', '[X]')

    # 转换 HTML 标签
    html_text = re.sub(r'<strong class="text-warning">(.*?)</strong>',
                       r'<font color="#f59e0b"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-danger">(.*?)</strong>',
                       r'<font color="#dc2626"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-success">(.*?)</strong>',
                       r'<font color="#10b981"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong class="text-info">(.*?)</strong>',
                       r'<font color="#0ea5e9"><b>\1</b></font>', html_text, flags=re.DOTALL)
    html_text = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', html_text, flags=re.DOTALL)

    html_text = html_text.replace('<br><br>', '|||PARA|||').replace('<br>', '<br/>')
    html_text = html_text.replace('<ul class="mb-0 mt-2">', '').replace('<ul>', '').replace('</ul>', '')
    html_text = re.sub(r'<li>(.*?)</li>', r'• \1<br/>', html_text, flags=re.DOTALL)

    # 先处理带 class 的 div 标签，再移除 class 属性
    html_text = html_text.replace('<div class="small text-muted mt-2">', '<font size="9" color="#6b7280">')
    html_text = re.sub(r' class="[^"]*"', '', html_text)

    paragraphs = html_text.split('|||PARA|||')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 移除空白的 div 标签
        para = re.sub(r'<div[^>]*>\s*</div>', '', para)
        # 处理剩余的 div 标签（转换为 font 或直接移除）
        para = para.replace('<div>', '<font size="9" color="#6b7280">')
        para = para.replace('</div>', '</font>')

        try:
            p = doc.add_paragraph()
            parse_paragraph_content(p, para)
        except Exception as e:
            print(f"Error rendering conclusion paragraph: {e}")
            plain_text = re.sub(r'<[^>]+>', '', para)
            doc.add_paragraph(plain_text)


def parse_paragraph_content(paragraph, html_text):
    """解析段落内容"""
    # 匹配模式（按优先级排序）
    pattern_colored_bold = r'<font color="([^"]+)"><b>(.*?)</b></font>'
    pattern_font = r'<font[^>]*size="(\d+)"[^>]*color="([^"]+)"[^>]*>(.*?)</font>'
    pattern_bold = r'<b>(.*?)</b>'

    html_text = html_text.replace('<br/>', '\n')

    pos = 0
    while pos < len(html_text):
        # 查找所有可能的匹配
        match_colored_bold = re.search(pattern_colored_bold, html_text[pos:], re.DOTALL)
        match_font = re.search(pattern_font, html_text[pos:], re.DOTALL)
        match_bold = re.search(pattern_bold, html_text[pos:], re.DOTALL)

        # 找出最近的匹配
        matches = []
        if match_colored_bold:
            matches.append(('colored_bold', match_colored_bold))
        if match_font:
            matches.append(('font', match_font))
        if match_bold:
            matches.append(('bold', match_bold))

        if not matches:
            # 没有更多标签，添加剩余文字
            remaining = html_text[pos:]
            if remaining:
                paragraph.add_run(remaining)
            break

        # 选择最近的匹配
        matches.sort(key=lambda x: x[1].start())
        match_type, next_match = matches[0]

        # 添加匹配前的普通文字
        if next_match.start() > 0:
            plain_text = html_text[pos:pos + next_match.start()]
            if plain_text:
                paragraph.add_run(plain_text)

        # 根据类型处理匹配
        if match_type == 'colored_bold':
            color = next_match.group(1)
            text = next_match.group(2)
            add_colored_text(paragraph, text, color=color, bold=True)
        elif match_type == 'font':
            size = int(next_match.group(1))
            color = next_match.group(2)
            text = next_match.group(3)
            add_colored_text(paragraph, text, color=color, bold=False, size=size)
        else:  # bold
            text = next_match.group(1)
            run = paragraph.add_run(text)
            run.bold = True

        pos = pos + next_match.end()


def generate_word_report_v2(data, output_file):
    """
    生成与 PDF 完全一致的 Word 报告
    """
    doc = Document()

    # 设定默认字体
    doc.styles['Normal'].font.name = 'Microsoft JhengHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # ========== 封面页 ==========
    # 只添加少量空行，讓標題從頂部開始
    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph('RELIABILITY TEST REPORT')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(32)
    title.runs[0].font.color.rgb = RGBColor(26, 84, 144)
    title.runs[0].font.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph('可靠度測試報告')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)
    subtitle.runs[0].font.color.rgb = RGBColor(55, 65, 81)

    # 減少空行以節省空間
    for _ in range(2):
        doc.add_paragraph()

    # 封面资讯表格
    cover_table = doc.add_table(rows=5, cols=2)

    # 设置列宽（匹配 PDF: 3", 3"）
    cover_table.columns[0].width = Inches(3.0)
    cover_table.columns[1].width = Inches(3.0)

    analysis_mode = data.get('analysis_mode', 'N/A').upper()

    cover_data = [
        ('Report Date / 報告日期', datetime.now().strftime('%Y-%m-%d')),
        ('Analysis Mode / 分析模式', analysis_mode),
        ('', ''),
        ('Prepared by / 報告產生', 'Reliability Analysis Tool'),
        ('', 'AFR/MTBF Analysis System')
    ]

    for i, (label, value) in enumerate(cover_data):
        cells = cover_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value

        # 设置单元格宽度
        cells[0].width = cover_table.columns[0].width
        cells[1].width = cover_table.columns[1].width

        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Microsoft JhengHei'
                    run.font.size = Pt(12)

        set_cell_background(cells[0], 'f9fafb')
        set_cell_background(cells[1], 'f9fafb')

    set_table_borders(cover_table, '2563eb', '18')

    # 分頁，讓 REPORT SUMMARY 從第二頁頂部開始
    doc.add_page_break()

    # ========== REPORT SUMMARY ==========
    summary_title = doc.add_heading('REPORT SUMMARY / 報告摘要', level=0)
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_title.runs[0].font.size = Pt(24)
    summary_title.runs[0].font.color.rgb = RGBColor(26, 84, 144)

    doc.add_paragraph()

    test_data = data.get('test_data', {})
    summary_table = doc.add_table(rows=4, cols=2)

    # 设置列宽（匹配 PDF: 2.5", 4"）
    summary_table.columns[0].width = Inches(2.5)
    summary_table.columns[1].width = Inches(4.0)

    summary_data = [
        ('Report Date / 報告日期:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('Analysis Mode / 分析模式:', data.get('analysis_mode', 'N/A').upper()),
        ('Test Duration / 測試時長:', f"{test_data.get('t_test', 'N/A')} hours"),
        ('Mission Time / 任務時間:', f"{test_data.get('mission_years', 'N/A')} years")
    ]

    for i, (label, value) in enumerate(summary_data):
        cells = summary_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value

        # 设置单元格宽度
        cells[0].width = summary_table.columns[0].width
        cells[1].width = summary_table.columns[1].width

        set_cell_background(cells[0], 'f3f4f6')
        set_cell_background(cells[1], 'f3f4f6')

        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft JhengHei'
                    run.font.size = Pt(10)

    set_table_borders(summary_table, 'd1d5db', '6')

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 1. Acceleration Factor Parameters ==========
    heading1 = doc.add_heading('1. Acceleration Factor Parameters / 加速因子參數', level=1)
    heading1.runs[0].font.size = Pt(16)
    heading1.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    af_params = data.get('af_params', {})
    af_param_table_data = [['Parameter / 參數', 'Use Condition / 使用條件', 'Test Condition / 測試條件', 'Model / 模型']]

    # 温度
    if af_params.get('enable_temp', True):
        af_param_table_data.append([
            'Temperature / 溫度',
            f"{af_params.get('t_use', '')}°C",
            f"{af_params.get('t_alt', '')}°C",
            f"Arrhenius (Ea={af_params.get('ea', '')} eV)"
        ])

    # 湿度
    if af_params.get('enable_hum', True):
        af_param_table_data.append([
            'Humidity / 濕度',
            f"{af_params.get('rh_use', '')}%",
            f"{af_params.get('rh_alt', '')}%",
            f"Peck (n={af_params.get('n_hum', '')})"
        ])

    # 电压
    if af_params.get('enable_voltage', False):
        af_param_table_data.append([
            'Voltage / 電壓',
            f"{af_params.get('v_use', '')}V",
            f"{af_params.get('v_alt', '')}V",
            f"IPL (β={af_params.get('beta_v', '')})"
        ])

    # 热循环
    if af_params.get('enable_tc', False):
        af_param_table_data.append([
            'Thermal Cycling / 熱循環',
            f"ΔT={af_params.get('dt_use', '')}°C, f={af_params.get('f_use', '')}/hr",
            f"ΔT={af_params.get('dt_alt', '')}°C, f={af_params.get('f_alt', '')}/hr",
            f"Coffin-Manson"
        ])

    # 振动
    if af_params.get('enable_vib', False):
        af_param_table_data.append([
            'Vibration / 機械振動',
            f"{af_params.get('g_use', '')}G",
            f"{af_params.get('g_alt', '')}G",
            f"IPL (n={af_params.get('n_vib', '')})"
        ])

    # 创建表格
    af_param_table = doc.add_table(rows=len(af_param_table_data), cols=4)

    # 设置列宽（匹配 PDF: 1.5", 1.8", 1.8", 2.4"）
    af_param_table.columns[0].width = Inches(1.5)
    af_param_table.columns[1].width = Inches(1.8)
    af_param_table.columns[2].width = Inches(1.8)
    af_param_table.columns[3].width = Inches(2.4)

    for i, row_data in enumerate(af_param_table_data):
        cells = af_param_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

            # 设置单元格自动换行
            cells[j].width = af_param_table.columns[j].width

            if i == 0:  # 表头
                set_cell_background(cells[j], '2563eb')
                cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[j].paragraphs[0].runs[0].bold = True
                cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            else:
                set_cell_background(cells[j], 'f9fafb')
                cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    set_table_borders(af_param_table, 'd1d5db', '6')

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 2. Acceleration Factor Results ==========
    heading2 = doc.add_heading('2. Acceleration Factor Results / 加速因子結果', level=1)
    heading2.runs[0].font.size = Pt(16)
    heading2.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    af_result = data.get('results', {}).get('af_result', {})
    af_result_table_data = [['Factor / 因子', 'Value / 數值']]

    if af_params.get('enable_temp', True):
        af_result_table_data.append(['AF (Temperature) / 溫度', format_af_value(af_result.get('af_t', 'N/A'))])
    if af_params.get('enable_hum', True):
        af_result_table_data.append(['AF (Humidity) / 濕度', format_af_value(af_result.get('af_rh', 'N/A'))])
    if af_params.get('enable_voltage', False):
        af_result_table_data.append(['AF (Voltage) / 電壓', format_af_value(af_result.get('af_v', 'N/A'))])

    # 空行
    af_result_table_data.append(['', ''])
    # 总计
    af_result_table_data.append(['AF Total / 總加速因子', format_af_value(af_result.get('af_total', 'N/A'))])

    af_result_table = doc.add_table(rows=len(af_result_table_data), cols=2)

    # 设置列宽（匹配 PDF: 4", 3.5"）
    af_result_table.columns[0].width = Inches(4.0)
    af_result_table.columns[1].width = Inches(3.5)

    for i, row_data in enumerate(af_result_table_data):
        cells = af_result_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

            # 设置单元格自动换行
            cells[j].width = af_result_table.columns[j].width

            if i == 0:  # 表头
                set_cell_background(cells[j], '2563eb')
                cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[j].paragraphs[0].runs[0].bold = True
                cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            elif i == len(af_result_table_data) - 1:  # 总计行
                set_cell_background(cells[j], 'dbeafe')
                cells[j].paragraphs[0].runs[0].bold = True
                cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            else:
                set_cell_background(cells[j], 'ffffff')
                cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    set_table_borders(af_result_table, 'd1d5db', '6')

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 3. Test Data ==========
    heading3 = doc.add_heading('3. Test Data / 測試數據', level=1)
    heading3.runs[0].font.size = Pt(16)
    heading3.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    test_table_data = [
        ['Test Time / 測試時間:', f"{test_data.get('t_test', 'N/A')} hours"],
        ['Number of Samples / 樣本數:', test_data.get('n_samples', 'N/A')],
        ['Failures / 失效數:', test_data.get('failures', '').count(',') + 1 if test_data.get('failures', '') else '0'],
        ['Confidence Level / 信心水準:', f"{float(test_data.get('cl', 0.6)) * 100:.0f}%"],
        ['Mission Time / 任務時間:', f"{test_data.get('mission_years', 'N/A')} years"],
        ['Equivalent Field Time / 等效現場時間:', f"{float(test_data.get('t_test', 0)) * float(af_result.get('af_total', 1)):,.2f} hours (~ {float(test_data.get('t_test', 0)) * float(af_result.get('af_total', 1)) / 8760:.2f} years)"]
    ]

    test_table = doc.add_table(rows=len(test_table_data), cols=2)

    # 设置列宽（匹配 PDF: 3", 4.5"）
    test_table.columns[0].width = Inches(3.0)
    test_table.columns[1].width = Inches(4.5)

    for i, row_data in enumerate(test_table_data):
        cells = test_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[1].text = str(row_data[1])

        # 设置单元格自动换行
        cells[0].width = test_table.columns[0].width
        cells[1].width = test_table.columns[1].width

        set_cell_background(cells[0], 'f3f4f6')
        set_cell_background(cells[1], 'f3f4f6')

        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # 最后一行高亮
        if i == len(test_table_data) - 1:
            set_cell_background(cells[0], 'dbeafe')
            set_cell_background(cells[1], 'dbeafe')

    set_table_borders(test_table, 'd1d5db', '6')

    doc.add_page_break()

    # ========== 4. Reliability Analysis Results ==========
    heading4 = doc.add_heading('4. Reliability Analysis Results / 可靠度分析結果', level=1)
    heading4.runs[0].font.size = Pt(16)
    heading4.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    reliability_result = data.get('results', {}).get('reliability_result', {})

    # Weibull 模式
    if 'weibull' in reliability_result:
        weibull_rel = reliability_result['weibull']

        weibull_title = doc.add_paragraph('Weibull Analysis Results:')
        weibull_title.runs[0].bold = True
        weibull_title.runs[0].font.size = Pt(12)

        doc.add_paragraph()

        weibull_table_data = [['Metric / 指標', 'Value / 數值']]
        weibull_table_data.append(['MTTF (Lower Limit) / 平均壽命下限:', f"{weibull_rel.get('mttf_use', 'N/A'):,.2f} hours"])
        weibull_table_data.append(['Max Failure Rate / 最大失效率:', '< N/A FITs'])
        weibull_table_data.append(['Reliability (2 Years) / 可靠度:', f"{weibull_rel.get('r_mission', 0) * 100:.2f}%"])
        weibull_table_data.append(['', ''])
        weibull_table_data.append(['Confidence Level / 信心水準:', f"{float(test_data.get('cl', 0.6)) * 100:.0f}%"])

        weibull_table = doc.add_table(rows=len(weibull_table_data), cols=2)

        # 设置列宽（匹配 PDF: 4", 3.5"）
        weibull_table.columns[0].width = Inches(4.0)
        weibull_table.columns[1].width = Inches(3.5)

        for i, row_data in enumerate(weibull_table_data):
            cells = weibull_table.rows[i].cells
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]

            # 设置单元格宽度
            cells[0].width = weibull_table.columns[0].width
            cells[1].width = weibull_table.columns[1].width

            if i == 0:
                set_cell_background(cells[0], '10b981')
                set_cell_background(cells[1], '10b981')
                cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[0].paragraphs[0].runs[0].bold = True
                cells[1].paragraphs[0].runs[0].bold = True
            else:
                set_cell_background(cells[0], 'ffffff')
                set_cell_background(cells[1], 'ffffff')

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        set_table_borders(weibull_table, 'd1d5db', '6')

    # Zero-failure 模式
    if 'zero_failure' in reliability_result:
        zf_rel = reliability_result['zero_failure']

        doc.add_paragraph()

        zf_title = doc.add_paragraph('Zero-Failure Analysis Results:')
        zf_title.runs[0].bold = True
        zf_title.runs[0].font.size = Pt(12)

        doc.add_paragraph()

        zf_table_data = [['Metric / 指標', 'Value / 數值']]
        zf_table_data.append(['MTTF (Lower Limit) / 平均壽命下限:', f"{zf_rel.get('mttf_use_lower', 'N/A'):,.2f} hours"])
        zf_table_data.append(['Max Failure Rate / 最大失效率:', f"< {zf_rel.get('lambda_use_upper', 'N/A'):,.2f} FITs"])
        zf_table_data.append(['Reliability (2 Years) / 可靠度:', f"{zf_rel.get('r_mission', 0) * 100:.5f}%"])
        zf_table_data.append(['', ''])
        zf_table_data.append(['Confidence Level / 信心水準:', f"{float(test_data.get('cl', 0.6)) * 100:.0f}%"])

        zf_table = doc.add_table(rows=len(zf_table_data), cols=2)

        # 设置列宽（匹配 PDF: 4", 3.5"）
        zf_table.columns[0].width = Inches(4.0)
        zf_table.columns[1].width = Inches(3.5)

        for i, row_data in enumerate(zf_table_data):
            cells = zf_table.rows[i].cells
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]

            # 设置单元格宽度
            cells[0].width = zf_table.columns[0].width
            cells[1].width = zf_table.columns[1].width

            if i == 0:
                set_cell_background(cells[0], '10b981')
                set_cell_background(cells[1], '10b981')
                cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cells[0].paragraphs[0].runs[0].bold = True
                cells[1].paragraphs[0].runs[0].bold = True
            else:
                set_cell_background(cells[0], 'ffffff')
                set_cell_background(cells[1], 'ffffff')

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        set_table_borders(zf_table, 'd1d5db', '6')

    # ========== 5. Charts (如果有) ==========
    charts_added = False
    try:
        from chart_generator import generate_all_charts

        analysis_mode = data.get('analysis_mode', 'weibull')
        weibull_result = data.get('results', {}).get('weibull_result')

        if weibull_result:
            charts = generate_all_charts(
                analysis_mode=analysis_mode,
                weibull_result=weibull_result if analysis_mode == 'weibull' else None,
                reliability_result=reliability_result,
                max_time=50000
            )

            valid_charts = {k: v for k, v in charts.items() if v}

            if valid_charts:
                doc.add_page_break()
                heading5 = doc.add_heading('5. Analysis Charts / 分析圖表', level=1)
                heading5.runs[0].font.size = Pt(16)
                heading5.runs[0].font.color.rgb = RGBColor(37, 99, 235)

                doc.add_paragraph()

                chart_titles = {
                    'reliability': 'Reliability Function R(t) / 可靠度曲線',
                    'failure_rate': 'Failure Rate λ(t) / 失效率曲線',
                    'pdf': 'Probability Density Function f(t) / 機率密度函數'
                }

                for i, (chart_name, chart_buffer) in enumerate(valid_charts.items(), start=1):
                    try:
                        chart_buffer.seek(0)

                        subtitle = doc.add_paragraph(f'5.{i} {chart_titles.get(chart_name, chart_name)}')
                        subtitle.runs[0].bold = True
                        subtitle.runs[0].font.size = Pt(11)

                        doc.add_paragraph()

                        doc.add_picture(chart_buffer, width=Inches(6.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        doc.add_paragraph()
                        charts_added = True
                    except Exception as e:
                        print(f"Error adding chart {chart_name}: {e}")

    except Exception as e:
        print(f"Charts not available: {e}")

    # ========== 6. Conclusion ==========
    doc.add_page_break()
    section_num = 6 if charts_added else 5
    heading_conclusion = doc.add_heading(f'{section_num}. Conclusion / 分析結論', level=1)
    heading_conclusion.runs[0].font.size = Pt(16)
    heading_conclusion.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    conclusion_html = data.get('conclusion', '')
    if conclusion_html and conclusion_html.strip():
        parse_html_conclusion(doc, conclusion_html)
    else:
        doc.add_paragraph('No conclusion available.')

    # 保存文件
    doc.save(output_file)


def generate_report_from_request_v2(data):
    """从 Flask 请求生成 Word 报告 V2"""
    try:
        buffer = io.BytesIO()
        generate_word_report_v2(data, buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error in generate_report_from_request_v2: {e}")
        import traceback
        traceback.print_exc()
        raise
